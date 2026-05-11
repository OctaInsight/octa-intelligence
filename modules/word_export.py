"""
Octa Intelligence — Word Document Export
Generates formatted Word documents for all analysis rounds.
"""
import io
from datetime import date


def _setup_doc():
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3)
        section.right_margin  = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    return doc


def _heading(doc, text, level=1, color_hex="1B2A4A"):
    from docx.shared import Pt, RGBColor
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(
            int(color_hex[0:2],16),
            int(color_hex[2:4],16),
            int(color_hex[4:6],16)
        )
    return h


def _bullet(doc, text, style_name="List Bullet"):
    try:
        p = doc.add_paragraph(text, style=style_name)
    except Exception:
        p = doc.add_paragraph(f"• {text}")
    return p


def export_round1_docx(analysis: dict, proposal_id: str,
                        acronym: str) -> bytes:
    """Export Round 1 Call Intelligence Analysis as Word document."""
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = _setup_doc()

    # Title
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run(f"Call Intelligence Analysis")
    run.bold = True; run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1B,0x2A,0x4A)

    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = t2.add_run(f"Proposal: {acronym} ({proposal_id})")
    run2.font.size = Pt(12); run2.font.color.rgb = RGBColor(0x88,0x99,0xB0)

    t3 = doc.add_paragraph()
    t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t3.add_run(f"Generated: {date.today().isoformat()}").font.color.rgb = RGBColor(0x88,0x99,0xB0)

    doc.add_page_break()

    # 1. Objectives
    _heading(doc, "1. Call Objectives")
    for obj in analysis.get("call_objectives",[]):
        _heading(doc, f"1.{obj.get('id','')} — {obj.get('title','')}", level=2)
        doc.add_paragraph(obj.get("description",""))
        if obj.get("evidence"):
            p = doc.add_paragraph()
            run = p.add_run(f"Evidence: {obj['evidence']}")
            run.italic = True; run.font.color.rgb = RGBColor(0x88,0x99,0xB0)

    # 2. Expected Outcomes
    _heading(doc, "2. Expected Outcomes")
    for oc in analysis.get("expected_outcomes",[]):
        _bullet(doc, f"{oc.get('outcome','')} [{oc.get('timeframe','')}]")

    # 3. Expected Impacts
    _heading(doc, "3. Expected Impacts")
    for imp in analysis.get("expected_impacts",[]):
        _bullet(doc, f"[{imp.get('level','').upper()}] {imp.get('impact','')} — {imp.get('timeframe','')} term")

    # 4. Expected Outputs
    _heading(doc, "4. Expected Outputs")
    for out in analysis.get("expected_outputs",[]):
        _bullet(doc, f"[{out.get('type','').upper()}] {out.get('output','')}")

    # 5. Recommended KPIs
    _heading(doc, "5. Recommended KPIs")
    for kpi in analysis.get("recommended_kpis",[]):
        _bullet(doc, f"{kpi.get('kpi','')} — Unit: {kpi.get('unit','')} [{kpi.get('source','')}]")

    # 6. Policy Framing
    _heading(doc, "6. Policy Framing")
    pf = analysis.get("policy_framing",{})
    if pf.get("policy_narrative"):
        doc.add_paragraph(pf["policy_narrative"])
    _heading(doc, "Primary Policies", level=2)
    for pol in pf.get("primary_policies",[]):
        _heading(doc, f"{pol.get('policy','')} (Relevance: {pol.get('relevance_score',0)}/10)", level=3)
        doc.add_paragraph(pol.get("how_call_references",""))
        if pol.get("key_targets_to_mention"):
            doc.add_paragraph("Key targets to mention:")
            for t in pol["key_targets_to_mention"]:
                _bullet(doc, t)

    # 7. Master Keywords
    _heading(doc, "7. Critical Keywords")
    critical = [k for k in analysis.get("master_keywords",[]) if k.get("importance")=="critical"]
    high     = [k for k in analysis.get("master_keywords",[]) if k.get("importance")=="high"]
    medium   = [k for k in analysis.get("master_keywords",[]) if k.get("importance")=="medium"]
    if critical:
        _heading(doc, "Critical (must appear):", level=2)
        doc.add_paragraph(" · ".join(k.get("keyword","") for k in critical))
    if high:
        _heading(doc, "High importance:", level=2)
        doc.add_paragraph(" · ".join(k.get("keyword","") for k in high))
    if medium:
        _heading(doc, "Medium importance:", level=2)
        doc.add_paragraph(" · ".join(k.get("keyword","") for k in medium))

    # 8. Synergy Initiatives
    _heading(doc, "8. Synergy & Complementarity")
    for syn in analysis.get("synergy_initiatives",[]):
        _heading(doc, syn.get("name",""), level=2)
        doc.add_paragraph(syn.get("description",""))
        if syn.get("how_to_build_synergy"):
            p = doc.add_paragraph()
            p.add_run("How to build synergy: ").bold = True
            p.add_run(syn["how_to_build_synergy"])

    # 9. Hidden Messages
    _heading(doc, "9. Hidden Messages & Implicit Expectations")
    doc.add_paragraph(analysis.get("hidden_messages",""))

    # 10. Expected Partner Profiles
    _heading(doc, "10. Expected Partner Profiles")
    for prof in analysis.get("expected_partner_profiles",[]):
        _heading(doc, prof.get("profile_type",""), level=2)
        doc.add_paragraph(prof.get("description",""))
        if prof.get("is_mandatory"):
            p = doc.add_paragraph()
            p.add_run("⚠ MANDATORY profile").bold = True

    # 11. Consortium Positioning
    _heading(doc, "11. Our Consortium Positioning")
    doc.add_paragraph(analysis.get("consortium_positioning",""))

    # 12. Strategic Recommendations
    _heading(doc, "12. Strategic Recommendations")
    for rec in sorted(analysis.get("strategic_recommendations",[]),
                       key=lambda x: x.get("priority",99)):
        _heading(doc, f"Priority {rec.get('priority','')}: {rec.get('recommendation','')}", level=2)
        doc.add_paragraph(rec.get("rationale",""))
        if rec.get("action"):
            p = doc.add_paragraph()
            p.add_run("Action: ").bold = True
            p.add_run(rec["action"])

    doc.add_paragraph(f"\n\nGenerated by Octa Intelligence · {date.today().isoformat()}")

    buf = io.BytesIO()
    doc.save(buf); buf.seek(0)
    return buf.read()


def export_architecture_docx(architecture: dict, analysis: dict,
                              proposal_id: str, acronym: str) -> bytes:
    """
    Export Round 2b proposal architecture as Word document.
    Red-text guidance appears under each section heading.
    """
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml    import OxmlElement

    RED   = RGBColor(0xC0,0x00,0x00)
    BLUE  = RGBColor(0x00,0x44,0x88)
    GREEN = RGBColor(0x00,0x66,0x33)
    AMBER = RGBColor(0xCC,0x66,0x00)

    doc = _setup_doc()
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # Title page
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run("Proposal Writing Architecture")
    run.bold = True; run.font.size = Pt(20); run.font.color.rgb = RGBColor(0x1B,0x2A,0x4A)

    doc.add_paragraph()
    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t2.add_run(f"{acronym} ({proposal_id})\n{architecture.get('programme','')}")

    doc.add_paragraph()
    t3 = doc.add_paragraph()
    t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t3.add_run(f"Generated: {date.today().isoformat()}")

    # Legend
    doc.add_page_break()
    _heading(doc, "How to Use This Document")
    legend_items = [
        (RED,   "Reviewer guidance — what evaluators expect in this section"),
        (BLUE,  "Policy connections — which EU policies to reference here"),
        (GREEN, "Keywords — must appear in this section"),
        (AMBER, "Evidence & measures — specific data, KPIs or proof expected"),
    ]
    for color, desc in legend_items:
        p = doc.add_paragraph()
        run = p.add_run(f"■ {desc}")
        run.font.color.rgb = color
    doc.add_paragraph()

    if architecture.get("general_advice"):
        _heading(doc, "General Strategic Advice")
        doc.add_paragraph(architecture["general_advice"])

    if architecture.get("top_5_priorities"):
        _heading(doc, "Top 5 Priorities for This Proposal")
        for item in architecture["top_5_priorities"]:
            _bullet(doc, item)

    doc.add_page_break()
    _heading(doc, "PROPOSAL STRUCTURE")

    def _safe_list(val):
        if not val: return []
        if isinstance(val, list): return [str(v) for v in val]
        return [str(val)]

    def _safe_pol(pol):
        if isinstance(pol, dict):
            text = pol.get("policy","") or pol.get("name","") or str(pol)
            how  = pol.get("how","") or pol.get("connection","")
            return f"{text} — {how}" if how else text
        return str(pol)

    sections = architecture.get("sections",[])
    for sec in sorted(sections, key=lambda x: x.get("order",0)):
        try:
            level = int(sec.get("level",1))
            title = str(sec.get("ai_title","") or sec.get("original_title","") or sec.get("title","Untitled"))
            _heading(doc, title, level=level)

            # Enhanced title note
            if sec.get("title_rationale"):
                p = doc.add_paragraph()
                run = p.add_run(f"[{sec['title_rationale']}]")
                run.italic = True
                run.font.color.rgb = RGBColor(0x88,0x99,0xB0)
                run.font.size = Pt(9)

            # ── 1. WHAT TO WRITE HERE (guidance paragraph) ──────────────────
            gp = str(sec.get("guidance_paragraph","")).strip()
            if gp:
                p   = doc.add_paragraph()
                run = p.add_run("WHAT TO WRITE HERE:")
                run.bold = True; run.font.color.rgb = RGBColor(0x00,0x66,0xAA)
                p2  = doc.add_paragraph(gp)
                p2.runs[0].font.color.rgb = RGBColor(0x00,0x44,0x88) if p2.runs else None

            # ── 2. GAP TO ADDRESS (red — from concept evaluation) ───────────
            gap = str(sec.get("gaps_to_address","")).strip()
            if gap:
                p   = doc.add_paragraph()
                run = p.add_run("GAP TO ADDRESS:")
                run.bold = True; run.font.color.rgb = RED
                gr  = p.add_run(f"  {gap}")
                gr.font.color.rgb = RED

            # ── 3. REVIEWER EXPECTS (red bullets) ───────────────────────────
            rg = _safe_list(sec.get("reviewer_guidance"))
            if rg:
                p   = doc.add_paragraph()
                run = p.add_run("REVIEWER EXPECTS:")
                run.bold = True; run.font.color.rgb = RED
                for point in rg:
                    bp = doc.add_paragraph(f"  * {point}", style="List Bullet")
                    for r in bp.runs: r.font.color.rgb = RED

            # ── 4. POLICY CONNECTIONS (blue) ─────────────────────────────────
            pc = sec.get("policy_connections")
            if pc:
                p   = doc.add_paragraph()
                run = p.add_run("POLICY CONNECTIONS:")
                run.bold = True; run.font.color.rgb = BLUE
                items = pc if isinstance(pc, list) else [pc]
                for pol in items:
                    bp = doc.add_paragraph(f"  * {_safe_pol(pol)}", style="List Bullet")
                    for r in bp.runs: r.font.color.rgb = BLUE

            # ── 5. KEYWORDS (green) ───────────────────────────────────────────
            kws = _safe_list(sec.get("keywords_to_include"))
            if kws:
                p   = doc.add_paragraph()
                run = p.add_run("KEYWORDS: ")
                run.bold = True; run.font.color.rgb = GREEN
                kw  = p.add_run(" / ".join(kws))
                kw.font.color.rgb = GREEN

            # ── 6. EVIDENCE NEEDED (amber) ────────────────────────────────────
            ev = sec.get("measures_and_evidence","")
            if ev:
                p   = doc.add_paragraph()
                run = p.add_run("EVIDENCE NEEDED: ")
                run.bold = True; run.font.color.rgb = AMBER
                er  = p.add_run(str(ev))
                er.font.color.rgb = AMBER

            # ── 7. Word count + mistakes ──────────────────────────────────────
            wc = sec.get("word_count_guidance","")
            if wc:
                p   = doc.add_paragraph()
                run = p.add_run(f"Recommended length: {wc}")
                run.font.color.rgb = RGBColor(0x88,0x99,0xB0)
                run.font.size = Pt(9)

            cm = _safe_list(sec.get("common_mistakes"))
            if cm:
                p   = doc.add_paragraph()
                run = p.add_run("AVOID:")
                run.bold = True; run.font.color.rgb = RGBColor(0x99,0x33,0x00)
                for m in cm:
                    bp = doc.add_paragraph(f"  x {m}", style="List Bullet")
                    for r in bp.runs: r.font.color.rgb = RGBColor(0x99,0x33,0x00)

        except Exception as sec_err:
            doc.add_paragraph(f"[Section error: {sec_err}]")

        doc.add_paragraph()

    doc.add_paragraph(f"\nGenerated by Octa Intelligence · {date.today().isoformat()}")

    buf = io.BytesIO()
    doc.save(buf); buf.seek(0)
    return buf.read()


def export_round2a_docx(evaluation: dict, proposal_id: str, acronym: str) -> bytes:
    """Export Round 2a Concept Note Evaluation as a Word document."""
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    GREEN  = RGBColor(0x00, 0x66, 0x33)
    RED    = RGBColor(0xC0, 0x00, 0x00)
    AMBER  = RGBColor(0xCC, 0x66, 0x00)
    NAVY   = RGBColor(0x1B, 0x2A, 0x4A)
    TEAL   = RGBColor(0x00, 0x7B, 0x8A)
    GREY   = RGBColor(0x55, 0x66, 0x77)

    doc = _setup_doc()

    # Title
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run("Concept Note Evaluation Report")
    run.bold = True; run.font.size = Pt(18); run.font.color.rgb = NAVY

    doc.add_paragraph()
    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t2.add_run(f"Proposal: {acronym} ({proposal_id})").font.color.rgb = GREY

    t3 = doc.add_paragraph()
    t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t3.add_run(f"Generated: {date.today().isoformat()}").font.color.rgb = GREY
    doc.add_page_break()

    # Overall score
    overall = float(evaluation.get("overall_alignment_score", 0) or 0)
    verdict = evaluation.get("readiness_verdict", "")

    _heading(doc, "Overall Alignment Score")
    score_para = doc.add_paragraph()
    run_score  = score_para.add_run(f"{overall:.0f} / 100")
    run_score.bold = True; run_score.font.size = Pt(28)
    run_score.font.color.rgb = (GREEN if overall >= 75 else
                                (AMBER if overall >= 50 else RED))
    if verdict:
        vp = doc.add_paragraph()
        vp.add_run(f"Verdict: {verdict.replace('_', ' ').title()}").bold = True

    if evaluation.get("overall_comment"):
        doc.add_paragraph()
        doc.add_paragraph(evaluation["overall_comment"])

    doc.add_page_break()

    # Scores by dimension
    dims = evaluation.get("scores_by_dimension", {})
    if dims:
        _heading(doc, "Scores by Dimension")
        from docx.oxml.ns import qn
        from docx.oxml    import OxmlElement

        def _set_bg(cell, hex_color):
            tc   = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd  = OxmlElement("w:shd")
            shd.set(qn("w:fill"), hex_color)
            shd.set(qn("w:val"),  "clear")
            tcPr.append(shd)

        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        for i, h in enumerate(["Dimension", "Score", "Comment"]):
            cell = table.rows[0].cells[i]
            cell.text = h
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            _set_bg(cell, "1B2A4A")

        for dim, data in dims.items():
            score   = data.get("score", 0) if isinstance(data, dict) else int(data)
            comment = data.get("comment","") if isinstance(data, dict) else ""
            row     = table.add_row().cells
            row[0].text = dim.replace("_"," ").title()
            row[1].text = f"{score}/100"
            row[1].paragraphs[0].runs[0].bold = True
            row[1].paragraphs[0].runs[0].font.color.rgb = (
                GREEN if score >= 75 else (AMBER if score >= 50 else RED)
            )
            row[2].text = comment

    doc.add_paragraph()

    # Strengths
    strengths = evaluation.get("strengths", [])
    if strengths:
        _heading(doc, "✅ Strengths")
        for s in strengths:
            p   = doc.add_paragraph()
            dim = s.get("dimension","").replace("_"," ").title()
            run = p.add_run(f"[{dim}] ")
            run.bold = True; run.font.color.rgb = GREEN
            p.add_run(s.get("description",""))
            if s.get("evidence"):
                ep = doc.add_paragraph()
                ep.add_run(f"Evidence: {s['evidence']}").font.color.rgb = GREY
                ep.runs[0].italic = True
        doc.add_paragraph()

    # Gaps
    gaps = evaluation.get("gaps", [])
    if gaps:
        _heading(doc, "❌ Gaps to Address")
        SEV_COLORS = {"high": RED, "medium": AMBER, "low": GREY}
        for g in gaps:
            sev = g.get("severity","medium")
            p   = doc.add_paragraph()
            run = p.add_run(f"[{sev.title()}] {g.get('dimension','').replace('_',' ').title()}: ")
            run.bold = True; run.font.color.rgb = SEV_COLORS.get(sev, GREY)
            p.add_run(g.get("description",""))
            if g.get("suggestion"):
                sp = doc.add_paragraph()
                sp.add_run(f"💡 Suggestion: {g['suggestion']}").font.color.rgb = TEAL
        doc.add_paragraph()

    # Missing keywords
    kws = evaluation.get("missing_keywords", [])
    if kws:
        _heading(doc, "🔑 Missing Keywords")
        p = doc.add_paragraph()
        run = p.add_run("These critical keywords from the call are absent in your concept: ")
        run.font.color.rgb = GREY
        doc.add_paragraph(" · ".join(kws))
        doc.add_paragraph()

    # Recommendations
    recs = evaluation.get("recommendations", [])
    if recs:
        _heading(doc, "💡 Recommendations")
        for r in sorted(recs, key=lambda x: x.get("priority", 99)):
            _heading(doc, f"Priority {r.get('priority','')}: {r.get('action','')}", level=2)
            if r.get("rationale"):
                doc.add_paragraph(r["rationale"])
        doc.add_paragraph()

    doc.add_paragraph(f"\n\nGenerated by Octa Intelligence · {date.today().isoformat()}")

    buf = io.BytesIO()
    doc.save(buf); buf.seek(0)
    return buf.read()
